"""将项目核心代码以仿范文附录格式追加到 版3.docx 尾部"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

PROJECT = r'C:\Users\CoirRaincoat\PyCharmMiscProject\CSJ-MathModeling'
doc_path = os.path.join(PROJECT, '版3.docx')

doc = Document(doc_path)

# 确保中文字体
style = doc.styles['Normal']
style.font.size = Pt(10.5)

def add_code_section(doc, filename, description, code_text, max_lines=80):
    """添加一个代码小节：文件名标题 + 描述 + 代码框"""
    # Section title
    p = doc.add_paragraph()
    r = p.add_run(f'文件：{filename}')
    r.bold = True; r.font.size = Pt(11)
    p.space_before = Pt(12)
    
    if description:
        p2 = doc.add_paragraph()
        r2 = p2.add_run(f'介绍：{description}')
        r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(100,100,100)
    
    # Code in box — use a table with 1 cell, grey background
    lines = code_text.split('\n')
    # Limit to max_lines
    if len(lines) > max_lines:
        shown = lines[:max_lines]
    else:
        shown = lines
    
    # Create a single-cell table for the code box
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    
    # Clear default paragraph
    cell.paragraphs[0].clear()
    
    for line in shown:
        p_code = cell.add_paragraph()
        r_code = p_code.add_run(line)
        r_code.font.size = Pt(7)
        r_code.font.name = 'Consolas'
        p_code.space_before = Pt(0)
        p_code.space_after = Pt(0)
    
    # If truncated
    if len(lines) > max_lines:
        p_trunc = doc.add_paragraph()
        r_trunc = p_trunc.add_run(f'... (共 {len(lines)} 行，此处展示前 {max_lines} 行)')
        r_trunc.font.size = Pt(8); r_trunc.font.color.rgb = RGBColor(150,150,150)
    
    doc.add_paragraph()  # spacing


# ==================== 附录标题 ====================
p = doc.add_paragraph()
r = p.add_run('附录：核心代码')
r.bold = True; r.font.size = Pt(14)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# ==================== config.py ====================
with open(os.path.join(PROJECT, 'config.py'), 'r', encoding='utf-8') as f:
    code = f.read()
# Extract key constants section only (skip long docstring)
lines = code.split('\n')
key_start = 0
for i, line in enumerate(lines):
    if 'BASE_DIR' in line and 'os.path' in line:
        key_start = i
        break
key_lines = lines[:5] + lines[key_start:key_start+120]  # docstring header + first 120 lines of body
add_code_section(doc, 'config.py', '全局配置文件：路径匹配、餐次划分、菜品分类规则、营养标准、NPG配色',
                 '\n'.join(key_lines))

# ==================== data_loader.py ====================
with open(os.path.join(PROJECT, 'data_loader.py'), 'r', encoding='utf-8') as f:
    code = f.read()
lines = code.split('\n')
# Extract key methods: _load_data, _build_basket
key_parts = []
in_load = False; in_basket = False; line_count = 0
for line in lines:
    if 'def _load_data(self):' in line:
        in_load = True; key_parts.append('# --- 数据加载 ---')
        line_count = 0
    elif 'def _build_basket(self):' in line:
        in_load = False; in_basket = True; key_parts.append('\n# --- 购物篮构建 ---')
        line_count = 0
    elif in_load or in_basket:
        if line.strip().startswith('def ') and line_count > 3:
            in_load = False; in_basket = False
        else:
            key_parts.append(line); line_count += 1
add_code_section(doc, 'data_loader.py', '数据加载与预处理：多Sheet全量加载、数据清洗、购物篮构建',
                 '\n'.join(key_parts[:70]))

# ==================== utils.py ====================
with open(os.path.join(PROJECT, 'utils.py'), 'r', encoding='utf-8') as f:
    code = f.read()
lines = code.split('\n')
# Extract mape_score, check_nutrition_balance
key_lines = [l for l in lines if not l.strip().startswith('"""') and not l.strip().startswith('#') and l.strip()][:80]
add_code_section(doc, 'utils.py', '工具函数：MAPE计算、滞后特征、滑动窗口、营养均衡度、热量分解',
                 '\n'.join(key_lines[:60]))

# ==================== problem2_prediction.py ====================
with open(os.path.join(PROJECT, 'problem2_prediction.py'), 'r', encoding='utf-8') as f:
    code = f.read()
lines = code.split('\n')
# Extract _sarima_forecast, _predict_may_2025
key_parts = []
in_method = False
for line in lines:
    if 'def _sarima_forecast(self' in line:
        in_method = True; key_parts.append('# SARIMA 预测')
    elif 'def _predict_may_2025(self' in line:
        in_method = True; key_parts.append('\n# May2025 样本外预测')
    elif in_method:
        if line.strip().startswith('def ') and len(key_parts) > 2:
            in_method = False
        else:
            key_parts.append(line)
add_code_section(doc, 'problem2_prediction.py', '问题二：SARIMA预测与May2025外推（核心方法节选）',
                 '\n'.join(key_parts[:80]))

# ==================== problem3_optimization.py ====================
with open(os.path.join(PROJECT, 'problem3_optimization.py'), 'r', encoding='utf-8') as f:
    code = f.read()
lines = code.split('\n')
key_parts = []
in_opt = False
for line in lines:
    if 'def optimize_meal(self' in line:
        in_opt = True; key_parts.append('# MILP 优化核心函数')
    elif in_opt:
        if line.strip().startswith('def ') and len(key_parts) > 10:
            in_opt = False
        else:
            key_parts.append(line)
add_code_section(doc, 'problem3_optimization.py', '问题三：MILP午餐备菜优化（目标函数+约束+求解）',
                 '\n'.join(key_parts[:90]))

# ==================== problem4_combos.py ====================
with open(os.path.join(PROJECT, 'problem4_combos.py'), 'r', encoding='utf-8') as f:
    code = f.read()
lines = code.split('\n')
key_parts = []
in_score = False
for line in lines:
    if 'def _score_combo(self' in line:
        in_score = True; key_parts.append('# 套餐五维评分函数')
    elif in_score:
        if line.strip().startswith('def ') and len(key_parts) > 10:
            in_score = False
        else:
            key_parts.append(line)
add_code_section(doc, 'problem4_combos.py', '问题四：套餐五维评分函数与贪心搜索',
                 '\n'.join(key_parts[:80]))

# ==================== problem5_strategy.py ====================
with open(os.path.join(PROJECT, 'problem5_strategy.py'), 'r', encoding='utf-8') as f:
    code = f.read()
lines = code.split('\n')
key_parts = []
in_strat = False
for line in lines:
    if 'def _preparation_strategy(self' in line:
        in_strat = True; key_parts.append('# 备菜策略分析')
    elif in_strat:
        if line.strip().startswith('def ') and len(key_parts) > 10:
            in_strat = False
        else:
            key_parts.append(line)
add_code_section(doc, 'problem5_strategy.py', '问题五：五维度经营策略分析',
                 '\n'.join(key_parts[:60]))

# ==================== main.py ====================
with open(os.path.join(PROJECT, 'main.py'), 'r', encoding='utf-8') as f:
    code = f.read()
add_code_section(doc, 'main.py', '主入口：串联全部模块（支持--skip/--only命令行参数）',
                 code, max_lines=60)

# Save
doc.save(doc_path)
print(f'Saved with code appendix: {doc_path}')
