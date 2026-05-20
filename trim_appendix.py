"""精简 版3.docx 附录代码 — 定位所有代码表并替换为核心片段"""
from docx import Document
from docx.shared import Pt
import os

PROJECT = r'C:\Users\CoirRaincoat\PyCharmMiscProject\CSJ-MathModeling'
DOC_PATH = os.path.join(PROJECT, '版3.docx')

doc = Document(DOC_PATH)

code_files = ['config.py','data_loader.py','utils.py','problem1_analysis.py',
              'problem2_prediction.py','problem3_optimization.py','problem4_combos.py',
              'problem5_strategy.py','main.py']

# Find all tables in the appendix (last section of doc)
all_tables = doc.tables
# Appendix tables are at the end — code tables are single-row, single-cell
code_table_indices = []
for idx, t in enumerate(all_tables):
    if len(t.rows) == 1 and len(t.rows[0].cells) == 1:
        ct = t.rows[0].cells[0].text[:80]
        if any(kw in ct for kw in ['import ','def ','BASE_DIR','# ','plt.','pd.','class ','from ','"""']):
            code_table_indices.append(idx)

# If not enough, take last N tables
if len(code_table_indices) < len(code_files):
    code_table_indices = list(range(max(0, len(all_tables)-len(code_files)), len(all_tables)))

print(f'Code tables: {len(code_table_indices)} (files: {len(code_files)})')

def extract_core(filename, max_lines=35):
    fp = os.path.join(PROJECT, filename)
    if not os.path.exists(fp): return '# file not found'
    with open(fp, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    kw_map = {
        'config.py': ['COST_RATIO_BY_CATEGORY', 'COLORS =', 'NUTRITION_PER_MEAL', 'SAFETY_STOCK_FACTOR'],
        'data_loader.py': ['def _load_data', 'def _build_basket'],
        'utils.py': ['def mape_score', 'def check_nutrition_balance'],
        'problem1_analysis.py': ['def _sales_distribution_analysis', 'def _association_rule_mining'],
        'problem2_prediction.py': ['def _predict_may_2025', 'def _walk_forward_validation'],
        'problem3_optimization.py': ['def optimize_meal', 'PULP_CBC_CMD'],
        'problem4_combos.py': ['def _score_combo', 'def _greedy_search'],
        'problem5_strategy.py': ['def _plot_strategy_summary', 'def _preparation_strategy'],
        'main.py': ['def run_all', 'def should_run'],
    }
    
    keywords = kw_map.get(filename, ['def '])
    result, capture, indent0, lc = [], False, '', 0
    
    for line in lines:
        s = line.strip()
        if not capture:
            for kw in keywords:
                if kw in s:
                    capture = True
                    indent0 = line[:len(line)-len(line.lstrip())]
                    break
        if capture:
            if lc >= max_lines:
                result.append(f'# ... (共 {len(lines)} 行)'); break
            result.append(line.rstrip()); lc += 1
            if s.startswith('def ') and lc > 5 and line[:len(line)-len(line.lstrip())] == indent0:
                if not any(kw in s for kw in keywords):
                    break
    
    return '\n'.join(result) if result else f'# 未匹配, 共{len(lines)}行'

# Update tables
for i, tidx in enumerate(code_table_indices):
    if i >= len(code_files): break
    fname = code_files[i]
    code = extract_core(fname, max_lines=35)
    cell = all_tables[tidx].rows[0].cells[0]
    for p in cell.paragraphs: p.clear()
    p0 = cell.paragraphs[0]; p0.clear()
    for line in code.split('\n'):
        r = p0.add_run(line + '\n'); r.font.size = Pt(6.5); r.font.name = 'Consolas'
    print(f'  [{i}] {fname}: {len(code.splitlines())} lines')

doc.save(DOC_PATH)
print(f'Saved: {DOC_PATH}')
